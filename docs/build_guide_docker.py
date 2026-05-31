#!/usr/bin/env python3
"""Generate the Docker / Docker Compose installation guide (.docx).

Pedagogical handout for the Master 2 IA/Cybersecurite lab (JuiceLab).
Output is a true .docx binary produced with python-docx.

Run:
    python build_guide_docker.py [output_path]

Default output: Guide_Installation_Docker_JuiceLab.docx
"""

from __future__ import annotations

import logging
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# Palette (institutional, sober).
COLOR_TITLE = RGBColor(0x1F, 0x36, 0x4D)
COLOR_HEADING = RGBColor(0x2E, 0x4B, 0x73)
COLOR_CODE_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_NOTE = RGBColor(0x8A, 0x33, 0x00)
SHADE_CODE = "F2F2F2"
SHADE_HEADER = "2E4B73"
SHADE_ROW_ALT = "EDF1F6"

BASE_FONT = "Arial"
MONO_FONT = "Consolas"

# Elements that, per the CT_PPr schema, must appear AFTER w:shd.
_SHD_SUCCESSORS = (
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)
# w:pBdr precedes w:shd in the schema.
_PBDR_SUCCESSORS = ("w:shd",) + _SHD_SUCCESSORS


def _set_cell_background(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_paragraph_shading(paragraph, hex_color: str) -> None:
    """Apply a background shading to a whole paragraph (used for code blocks)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.insert_element_before(shd, *_SHD_SUCCESSORS)


def _set_paragraph_border(paragraph, hex_color: str = "D0D0D0") -> None:
    """Draw a thin box around a paragraph (code block frame)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        elm = OxmlElement(f"w:{edge}")
        elm.set(qn("w:val"), "single")
        elm.set(qn("w:sz"), "4")
        elm.set(qn("w:space"), "6")
        elm.set(qn("w:color"), hex_color)
        borders.append(elm)
    p_pr.insert_element_before(borders, *_PBDR_SUCCESSORS)


def add_heading1(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    return p


def add_heading2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_note(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run_label = p.add_run(f"{label} ")
    run_label.font.name = BASE_FONT
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    run_label.font.color.rgb = COLOR_NOTE
    run_body = p.add_run(text)
    run_body.font.name = BASE_FONT
    run_body.font.size = Pt(11)
    return p


def add_code_block(doc: Document, lines: list[str]):
    """Render a monospaced, shaded, framed code block.

    Each element of ``lines`` is one line of code.
    """
    p = doc.add_paragraph()
    _set_paragraph_border(p)
    _set_paragraph_shading(p, SHADE_CODE)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.2)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = MONO_FONT
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), MONO_FONT)
        r_fonts.set(qn("w:hAnsi"), MONO_FONT)
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_CODE_TEXT
        if idx < len(lines) - 1:
            run.add_break()
    return p


def add_comparison_table(doc: Document) -> None:
    """Two-method comparison table for the Linux section."""
    headers = ["Critère", "Méthode A (distribution)", "Méthode B (dépôt officiel)"]
    rows = [
        ["Paquet compose", "docker-compose-v2", "docker-compose-plugin"],
        ["Moteur", "docker.io (tiré en dépendance)", "docker-ce (à installer)"],
        ["Commande d'install", "Une seule commande", "Configuration dépôt puis install"],
        ["Version", "Plus ancienne, stable", "Dernière version stable upstream"],
        ["Mises à jour", "Via apt upgrade système", "Via apt upgrade (dépôt Docker)"],
        ["Usage recommandé", "Lab pédagogique, poste étudiant", "Production, besoin des dernières features"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    widths = [Cm(4.0), Cm(6.0), Cm(6.0)]

    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = widths[col]
        _set_cell_background(cell, SHADE_HEADER)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(text)
        run.font.name = BASE_FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r, row in enumerate(rows, start=1):
        for col, text in enumerate(row):
            cell = table.cell(r, col)
            cell.width = widths[col]
            if r % 2 == 0:
                _set_cell_background(cell, SHADE_ROW_ALT)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(text)
            run.font.name = BASE_FONT
            run.font.size = Pt(10)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Master 2 IA / Cybersécurité - Lab JuiceLab - Guide d'installation Docker")
    run.font.name = BASE_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _fix_zoom(doc: Document) -> None:
    """Ensure w:zoom carries a valid percent attribute (python-docx default omits it)."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")


def build(output_path: str) -> None:
    logger.info("Building document: %s", output_path)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Installation de Docker et Docker Compose")
    run.font.name = BASE_FONT
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Guide d'installation pour le lab JuiceLab - Windows, macOS, Linux")
    run.font.name = BASE_FONT
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_HEADING
    subtitle.paragraph_format.space_after = Pt(14)

    # 1. Objectif et prerequis.
    add_heading1(doc, "1. Objectif et prérequis")
    add_body(
        doc,
        "Ce guide installe Docker Engine et Docker Compose v2 (commande docker compose, "
        "avec un espace) sur les trois familles de systèmes. Le tooling sert de socle au "
        "lab JuiceLab. La commande historique docker-compose (avec un trait d'union) est "
        "la version v1, dépréciée, et n'est pas utilisée ici.",
    )
    add_bullet(doc, "Accès administrateur sur le poste (ou droits sudo sous Linux).")
    add_bullet(doc, "Virtualisation matérielle activée dans le BIOS/UEFI (Windows et certains hyperviseurs).")
    add_bullet(doc, "Connexion réseau pour télécharger les images et les paquets.")

    # 2. Windows.
    add_heading1(doc, "2. Windows 10 / 11")
    add_body(doc, "Ouvrir PowerShell en tant qu'administrateur, installer WSL2 puis redémarrer :")
    add_code_block(doc, ["wsl --install"])
    add_body(
        doc,
        "La commande active WSL2, la fonctionnalité Plateforme de machine virtuelle et "
        "installe Ubuntu par défaut. Le redémarrage est obligatoire pour finaliser.",
    )
    add_bullet(doc, "Télécharger Docker Desktop depuis docker.com (backend WSL2 par défaut).")
    add_bullet(doc, "Lancer Docker Desktop et attendre le statut running (icône baleine).")
    add_note(
        doc,
        "Erreur de virtualisation :",
        "activer Virtualization / SVM / VT-x dans le BIOS/UEFI, puis vérifier que les "
        "fonctionnalités Windows Plateforme de machine virtuelle et Sous-système Windows "
        "pour Linux sont bien cochées (Panneau de configuration, Activer ou désactiver des "
        "fonctionnalités Windows).",
    )

    # 3. macOS.
    add_heading1(doc, "3. macOS")
    add_bullet(doc, "Télécharger Docker Desktop pour Mac en choisissant la puce : Apple Silicon ou Intel.")
    add_bullet(doc, "Glisser Docker dans Applications, le lancer, attendre le statut running.")
    add_note(
        doc,
        "Attention :",
        "le seul piège courant est le choix de l'architecture. Un binaire Intel sur Apple "
        "Silicon fonctionne via Rosetta mais dégrade les performances ; préférer le binaire natif.",
    )

    # 4. Linux.
    add_heading1(doc, "4. Linux (Debian / Ubuntu)")
    add_body(
        doc,
        "Il existe deux méthodes d'installation distinctes et mutuellement exclusives. "
        "Le point essentiel est de choisir l'une OU l'autre, jamais un mélange des deux : "
        "le paquet de la distribution docker-compose-v2 et le plugin du dépôt officiel "
        "docker-compose-plugin entrent en conflit.",
    )

    add_heading2(doc, "4.1 Méthode A - paquets de la distribution (la plus simple)")
    add_body(
        doc,
        "Le paquet docker-compose-v2 provient des dépôts Ubuntu/Debian et tire docker.io "
        "(le moteur) comme dépendance. Une seule commande suffit.",
    )
    add_code_block(
        doc,
        [
            "sudo apt update",
            "sudo apt install -y docker-compose-v2",
            "sudo usermod -aG docker $USER",
            "newgrp docker   # ou se déconnecter puis se reconnecter",
            "docker compose version",
        ],
    )
    add_body(
        doc,
        "Avantage : intégration native avec les mises à jour de sécurité apt. Inconvénient : "
        "version plus ancienne que l'upstream. Pour un lab pédagogique, c'est le choix "
        "pragmatique, la fraîcheur de version n'ayant aucune importance ici.",
    )

    add_heading2(doc, "4.2 Méthode B - dépôt officiel Docker (la plus à jour)")
    add_body(
        doc,
        "Après avoir configuré le dépôt apt officiel (clé GPG plus fichier sources dans "
        "/etc/apt/sources.list.d/), installer le moteur et le plugin compose. Dans ce cas, "
        "on n'installe jamais docker-compose-v2 : le plugin compose arrive via "
        "docker-compose-plugin et docker compose fonctionne directement.",
    )
    add_code_block(
        doc,
        [
            "# configurer d'abord le dépôt officiel :",
            "# https://docs.docker.com/engine/install/ubuntu/",
            "sudo apt install -y docker-ce docker-ce-cli containerd.io \\",
            "    docker-buildx-plugin docker-compose-plugin",
            "sudo usermod -aG docker $USER",
            "newgrp docker",
            "docker compose version",
        ],
    )
    add_note(
        doc,
        "Piège à connaître :",
        "avant l'installation via le dépôt officiel, il faut purger les anciens paquets "
        "dont docker.io, docker-compose et docker-compose-v2, sinon conflit. Si un poste a "
        "déjà reçu la méthode A puis veut passer à la méthode B, nettoyer d'abord avec "
        "sudo apt remove docker-compose-v2 docker.io.",
    )
    add_comparison_table(doc)

    # 5. Verification.
    add_heading1(doc, "5. Vérification (commune à tous les systèmes)")
    add_body(doc, "Une fois Docker en place, valider l'installation :")
    add_code_block(
        doc,
        [
            "docker run --rm hello-world",
            "docker compose version",
        ],
    )
    add_body(
        doc,
        "La première commande télécharge une image de test, la lance, affiche un message "
        "de confirmation puis sort. La seconde confirme que le plugin compose v2 répond.",
    )

    # 6. Depannage.
    add_heading1(doc, "6. Dépannage")
    add_heading2(doc, "Permission denied au lancement de docker")
    add_body(
        doc,
        "L'utilisateur n'est pas dans le groupe docker. Exécuter sudo usermod -aG docker "
        "$USER puis newgrp docker (ou se reconnecter). Tant que la session n'a pas été "
        "rouverte, le groupe n'est pas appliqué.",
    )
    add_heading2(doc, "Conflit de paquets compose")
    add_body(
        doc,
        "Symptôme typique : docker compose introuvable après une install partielle, ou "
        "erreur de paquet en conflit. Cela vient du mélange méthode A et méthode B. "
        "Choisir une seule méthode et purger l'autre famille de paquets.",
    )
    add_heading2(doc, "Virtualisation désactivée (Windows)")
    add_body(
        doc,
        "Docker Desktop refuse de démarrer si la virtualisation est inactive. Activer "
        "VT-x / SVM dans le BIOS/UEFI et cocher les fonctionnalités Windows mentionnées "
        "en section 2.",
    )
    add_heading2(doc, "Poste verrouillé ou sans droits (contexte TD)")
    add_body(
        doc,
        "Si la DSI bloque la virtualisation, l'accès BIOS ou Docker Desktop, deux plans B "
        "évitent de bloquer le TD :",
    )
    add_bullet(doc, "Lancer JuiceLab directement en Python/Flask sans conteneur sur ces postes.")
    add_bullet(doc, "Utiliser Podman en mode rootless, qui ne nécessite pas Docker Desktop.")

    # 7. References.
    add_heading1(doc, "7. Références")
    add_bullet(doc, "Installation officielle de Docker Engine : https://docs.docker.com/engine/install/")
    add_bullet(doc, "Plugin Docker Compose : https://docs.docker.com/compose/install/linux/")
    add_bullet(doc, "Docker Desktop : https://www.docker.com/products/docker-desktop/")

    _fix_zoom(doc)
    doc.save(output_path)
    logger.info("Document saved: %s", output_path)


def main() -> None:
    output = sys.argv[1] if len(sys.argv) > 1 else "Guide_Installation_Docker_JuiceLab.docx"
    build(output)


if __name__ == "__main__":
    main()
