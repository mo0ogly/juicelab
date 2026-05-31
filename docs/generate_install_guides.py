#!/usr/bin/env python3
"""Genere les guides d'installation JuiceLab au format .docx (eleve + prof).

Produit deux documents Word distribuables, couvrant Windows / macOS / Linux,
avec des diagrammes Mermaid rendus en PNG et embarques :

    docs/GUIDE-INSTALL-ELEVE.docx
    docs/GUIDE-INSTALL-PROF.docx

Dependances :
    - python-docx        (pip install python-docx)
    - @mermaid-js/mermaid-cli  -> binaire `mmdc` (npm install -g @mermaid-js/mermaid-cli)
      + un navigateur Chromium/Chrome disponible (mmdc en a besoin pour le rendu).

Usage :
    python docs/generate_install_guides.py

Source de verite : les valeurs ci-dessous (REPO_URL, COHORT, DASHBOARD_PORT)
refletent docker/.env et scripts/install-student.{sh,ps1}. Si l'un change,
mettre a jour ici et relancer.
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor
except ImportError:
    sys.exit(
        "python-docx introuvable. Installer : pip install python-docx "
        "(ou via un venv : python3 -m venv .venv && .venv/bin/pip install python-docx)."
    )

# ---------------------------------------------------------------------------
# Constantes projet (alignees sur docker/.env + scripts/install-student.*)
# ---------------------------------------------------------------------------

REPO_URL = "https://github.com/mo0ogly/juicelab.git"
COHORT = "M2-IA-2026"
DASHBOARD_PORT = "5050"
SHOP_PORT = "3000"
EXAMPLE_IP = "192.168.1.10"

DOCS_DIR = Path(__file__).resolve().parent
ELEVE_OUT = DOCS_DIR / "GUIDE-INSTALL-ELEVE.docx"
PROF_OUT = DOCS_DIR / "GUIDE-INSTALL-PROF.docx"

# Palette
BLUE = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY_FILL = "F2F2F2"
NOTE_FILL = "FFF4D6"
HEAD_FILL = "D5E8F0"

# ---------------------------------------------------------------------------
# Rendu Mermaid -> PNG
# ---------------------------------------------------------------------------


def render_mermaid(source: str, out_png: Path, tmp: Path) -> bool:
    """Rend un diagramme Mermaid en PNG via mmdc. Retourne True si succes."""
    mmdc = shutil.which("mmdc")
    if not mmdc:
        print("  [warn] mmdc introuvable — diagramme saute. "
              "Installer : npm install -g @mermaid-js/mermaid-cli", file=sys.stderr)
        return False

    mmd = tmp / (out_png.stem + ".mmd")
    mmd.write_text(source.strip() + "\n", encoding="utf-8")

    # Config puppeteer : --no-sandbox indispensable en environnement CI / root.
    pcfg = tmp / "puppeteer.json"
    pcfg.write_text('{"args":["--no-sandbox","--disable-gpu"]}', encoding="utf-8")

    cmd = [
        mmdc, "-i", str(mmd), "-o", str(out_png),
        "-b", "white", "-s", "2", "-p", str(pcfg),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0 or not out_png.exists():
        print(f"  [warn] echec rendu {out_png.name} :\n{proc.stderr.strip()}",
              file=sys.stderr)
        return False
    return True


# ---------------------------------------------------------------------------
# Helpers python-docx
# ---------------------------------------------------------------------------


def setup_styles(doc: Document) -> None:
    """A4, police par defaut, styles de titres."""
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 18), (2, 14), (3, 12)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLUE if level == 1 else DARK

    # Le template python-docx ecrit <w:zoom w:val="bestFit"/> ; le schema OOXML
    # exige w:percent. On corrige pour une validation propre.
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")
        if zoom.get(qn("w:val")) is not None:
            del zoom.attrib[qn("w:val")]


# Elements de w:pPr qui, selon le schema OOXML (CT_PPr), suivent w:shd.
# insert_element_before place w:shd juste avant le premier present -> ordre valide.
_PPR_AFTER_SHD = (
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


def _shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.insert_element_before(shd, *_PPR_AFTER_SHD)


def title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    r = sub.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = DARK
    r.italic = True
    # filet bleu sous le titre
    bar = doc.add_paragraph()
    pPr = bar._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code_block(doc: Document, lines: list[str]) -> None:
    """Bloc de commandes : police monospace, fond gris clair."""
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _shade(p, GREY_FILL)
        run = p.add_run(ln if ln else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def note(doc: Document, text: str) -> None:
    """Encadre d'avertissement (fond jaune pale)."""
    p = doc.add_paragraph()
    _shade(p, NOTE_FILL)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = False


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths_mm: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _shade(hdr[i].paragraphs[0], HEAD_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    for col_idx, w in enumerate(widths_mm):
        for cell in table.columns[col_idx].cells:
            cell.width = Mm(w)


def add_diagram(doc: Document, png: Path, caption: str) -> None:
    if png and png.exists():
        doc.add_picture(str(png), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = DARK


# ---------------------------------------------------------------------------
# Diagrammes Mermaid
# ---------------------------------------------------------------------------

MERMAID_SCENARIO4 = """
flowchart TB
    subgraph prof["PC enseignant - LAN classe"]
        D["juicelab-dashboard<br/>:%PORT%"]
    end
    subgraph toi["Ton poste eleve"]
        B["Navigateur<br/>127.0.0.1:%SHOP%"]
        J["juice-shop<br/>:%SHOP%"]
        B --> J
    end
    J -. "POST events<br/>http://ip-prof:%PORT%" .-> D
""".replace("%PORT%", DASHBOARD_PORT).replace("%SHOP%", SHOP_PORT)

MERMAID_SEQUENCE = """
sequenceDiagram
    autonumber
    participant E as Eleve (navigateur)
    participant J as juice-shop :%SHOP%
    participant D as dashboard :%PORT% (toi)
    E->>J: joue le TD (hints, quiz, flags)
    J->>D: POST /api/sync (events)
    D->>D: SQLite : matrice cohorte
    Note over D: tu lis /dashboard?cohort=%COHORT% + token
    D-->>E: rien (l'eleve ne voit jamais le dashboard)
""".replace("%PORT%", DASHBOARD_PORT).replace("%SHOP%", SHOP_PORT).replace("%COHORT%", COHORT)


# ---------------------------------------------------------------------------
# Document ELEVE
# ---------------------------------------------------------------------------


def build_eleve(tmp: Path) -> None:
    doc = Document()
    setup_styles(doc)
    title_block(doc, "JuiceLab — Guide d'installation (eleve)",
                "OWASP Juice Shop + parcours pedagogique — Windows / macOS / Linux")

    doc.add_heading("1. Ce que tu vas installer", level=1)
    para(doc, "JuiceLab fait tourner OWASP Juice Shop (la plateforme du TD) sur "
              "ton propre poste. Pendant que tu joues, ton poste envoie ta "
              "progression au dashboard du professeur, qui suit toute la cohorte "
              "en temps reel. Tu n'ouvres jamais le dashboard toi-meme.")

    png4 = tmp / "scenario4.png"
    render_mermaid(MERMAID_SCENARIO4, png4, tmp)
    add_diagram(doc, png4, "Ton poste pousse ses events vers le dashboard du prof (scenario 4).")

    doc.add_heading("2. Ce que le prof te donne", level=1)
    para(doc, "Avant de commencer, recupere ces trois informations aupres du professeur :")
    add_table(
        doc,
        ["Information", "Exemple", "A quoi ca sert"],
        [
            ["IP du dashboard prof", EXAMPLE_IP, "Ou ton poste envoie ta progression"],
            ["Identifiant de cohorte", COHORT, "Te regroupe avec ta classe"],
            ["Ton label (prenom)", "amelie", "T'identifie dans la matrice du prof"],
        ],
        [55, 45, 70],
    )

    doc.add_heading("3. Prerequis", level=1)
    add_table(
        doc,
        ["Systeme", "A installer"],
        [
            ["Windows 10/11", "Docker Desktop (avec WSL2) + Git for Windows (fournit Git et OpenSSL)"],
            ["macOS", "Docker Desktop + Git (xcode-select --install)"],
            ["Linux", "Docker Engine + docker compose v2 + git + openssl (voir Annexe B)"],
        ],
        [45, 125],
    )
    note(doc, "Docker doit etre demarre (Docker Desktop ouvert) avant de lancer "
              "l'installation. Le premier build prend 5 a 8 minutes.")

    doc.add_heading("4. Installation — Windows (PowerShell)", level=1)
    para(doc, "Ouvre PowerShell, puis :")
    code_block(doc, [
        "git clone " + REPO_URL,
        "cd juicelab",
        "# Remplace l'IP par celle du prof, et amelie par ton prenom :",
        ".\\scripts\\install-student.ps1 -Dashboard " + EXAMPLE_IP +
        " -Label amelie -Cohort " + COHORT,
    ])
    note(doc, "Si PowerShell bloque le script : "
              "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass puis relance.")

    doc.add_heading("5. Installation — macOS / Linux (Terminal)", level=1)
    para(doc, "Ouvre un terminal, puis :")
    code_block(doc, [
        "git clone " + REPO_URL,
        "cd juicelab",
        "# Remplace l'IP par celle du prof, et amelie par ton prenom :",
        "./scripts/install-student.sh -d " + EXAMPLE_IP +
        " -l amelie -c " + COHORT,
    ])
    note(doc, "macOS (Apple Silicon M1/M2/M3) : aucune manipulation Rosetta requise, "
              "l'image se construit en natif arm64 ; le premier build peut juste etre "
              "un peu plus long. Docker Desktop doit etre ouvert avant de lancer.")

    doc.add_heading("6. Verifier que ca marche", level=1)
    numbered(doc, "Le script affiche « Installation OK » avec un recapitulatif.")
    numbered(doc, "Ouvre ton navigateur sur l'adresse ci-dessous :")
    code_block(doc, ["http://127.0.0.1:" + SHOP_PORT + "/#/juicelab"])
    note(doc, "Tu vois « Connecte-toi a Juice Shop » ? C'est normal avant le login. "
              "Va sur /#/register, cree un compte, connecte-toi, puis reviens ici — "
              "le panneau s'affiche automatiquement.")
    note(doc, "Le script affiche « !!! Dashboard prof injoignable » ? Normal si le prof "
              "n'a pas encore lance son dashboard ou que tu n'es pas sur le meme reseau. "
              "L'installation est reussie : les events partiront des que le dashboard sera disponible.")
    numbered(doc, "Tu dois voir le parcours TD JuiceLab. Le score-board OWASP est sur "
                  "http://127.0.0.1:" + SHOP_PORT + "/#/score-board.")
    para(doc, "Des que tu reveles un indice ou resous un challenge, ta progression "
              "remonte automatiquement chez le prof.")
    add_table(
        doc,
        ["Identifiant", "Origine", "Role"],
        [
            ["Label (-l fabrice)",
             "docker/.env, defini par le prof a l'install",
             "Identifie ton poste dans la matrice cohorte — fixe, independant de Juice Shop"],
            ["Email Juice Shop",
             "Compte cree sur /#/register",
             "Deverrouille le panneau JuiceLab — peut etre n'importe quelle adresse"],
        ],
        [40, 65, 65],
    )
    note(doc, "Le prof voit la colonne avec ton label (ex. « fabrice ») dans sa matrice. "
              "L'email Juice Shop n'est jamais affiche en TD standard. "
              "Deux options : (1) email fictif type fabrice@juicelab.local — format x@y.z requis, "
              "adresse non verifiee. (2) Login Google : le bouton fonctionne sur 127.0.0.1:3000 "
              "grace au proxy OWASP (local3000.owasp-juice.shop) qui redirige le callback OAuth "
              "vers localhost.")

    doc.add_heading("7. Depannage", level=1)
    add_table(
        doc,
        ["Symptome", "Cause probable / solution"],
        [
            ["Le script dit que Docker est introuvable",
             "Docker Desktop n'est pas demarre. Ouvre-le et relance."],
            ["« Dashboard prof injoignable »",
             "Verifie l'IP donnee par le prof, que tu es sur le meme reseau, "
             "et que le port " + DASHBOARD_PORT + " n'est pas bloque par un pare-feu."],
            ["La page 127.0.0.1:" + SHOP_PORT + " ne charge pas",
             "Attends la fin du build (5-8 min) puis rafraichis."],
            ["Ma progression n'apparait pas chez le prof",
             "Verifie que tu as bien passe -d / -Dashboard avec la bonne IP "
             "et la meme cohorte (" + COHORT + ") que la classe."],
        ],
        [55, 115],
    )

    doc.add_heading("8. Arreter / nettoyer", level=1)
    code_block(doc, [
        "cd docker",
        "docker compose --env-file .env down       # arret (garde les donnees)",
        "docker compose --env-file .env down -v    # arret + efface tout",
    ])

    annexes(doc)
    doc.save(str(ELEVE_OUT))
    print(f"  ecrit : {ELEVE_OUT}")


# ---------------------------------------------------------------------------
# Document PROF
# ---------------------------------------------------------------------------


def build_prof(tmp: Path) -> None:
    doc = Document()
    setup_styles(doc)
    title_block(doc, "JuiceLab — Guide d'installation (professeur)",
                "Dashboard de consolidation de cohorte — Windows / macOS / Linux")

    doc.add_heading("1. Ton role", level=1)
    para(doc, "Tu fais tourner uniquement le dashboard sur ton poste. Chaque eleve "
              "installe juice-shop sur le sien et le configure pour pousser sa "
              "progression vers ton dashboard. Tu obtiens une matrice cohorte "
              "(eleves x challenges) en temps reel.")

    pngseq = tmp / "sequence.png"
    render_mermaid(MERMAID_SEQUENCE, pngseq, tmp)
    add_diagram(doc, pngseq, "Flux des events : l'eleve joue, ton dashboard consolide.")

    png4 = tmp / "scenario4_prof.png"
    render_mermaid(MERMAID_SCENARIO4, png4, tmp)
    add_diagram(doc, png4, "Topologie scenario 4 : un dashboard (toi), N postes eleves.")

    doc.add_heading("2. Prerequis", level=1)
    add_table(
        doc,
        ["Systeme", "A installer"],
        [
            ["Windows 10/11", "Docker Desktop (avec WSL2) + Git for Windows"],
            ["macOS", "Docker Desktop + Git"],
            ["Linux", "Docker Engine + docker compose v2 + git + openssl (voir Annexe B)"],
        ],
        [45, 125],
    )
    note(doc, "Le dashboard est leger (~180 Mo, 8 Go RAM suffisent). Il doit etre "
              "joignable par les eleves : reseau LAN plat et port " + DASHBOARD_PORT +
              " ouvert dans le pare-feu.")

    doc.add_heading("3. Configurer les secrets", level=1)
    para(doc, "Le fichier docker/.env contient les jetons. Le script le cree a partir "
              "de docker/.env.example si besoin et genere les jetons manquants. "
              "Verifie ces valeurs (>= 16 caracteres) :")
    add_table(
        doc,
        ["Variable", "Role"],
        [
            ["DASHBOARD_TEACHER_TOKEN", "Jeton pour acceder a ton dashboard (a garder secret)"],
            ["DASHBOARD_PROOF_SECRET", "Signe les preuves de lab (verification des flags)"],
            ["DASHBOARD_CORS_ORIGINS", "Origines autorisees a poster ; defaut http://127.0.0.1:" + SHOP_PORT],
        ],
        [70, 100],
    )

    doc.add_heading("4. Lancer le dashboard", level=1)

    doc.add_heading("Windows (PowerShell)", level=2)
    code_block(doc, [
        "git clone " + REPO_URL,
        "cd juicelab",
        ".\\scripts\\install-student.ps1 -Server -Cohort " + COHORT,
    ])

    doc.add_heading("macOS / Linux (Terminal)", level=2)
    code_block(doc, [
        "git clone " + REPO_URL,
        "cd juicelab",
        "./scripts/install-student.sh --server -c " + COHORT,
    ])
    note(doc, "macOS : le script detecte ton IP LAN via ipconfig (pas besoin de "
              "hostname -I, propre a Linux). Sur Apple Silicon, le dashboard tourne "
              "en natif arm64.")

    doc.add_heading("5. Distribuer aux eleves", level=1)
    para(doc, "A la fin, le script affiche un bloc « A DISTRIBUER AUX ELEVES » "
              "avec ton IP LAN detectee et la commande exacte que chaque eleve doit "
              "lancer. Exemple :")
    code_block(doc, [
        "A DISTRIBUER AUX ELEVES (scenario 4) :",
        "  Cohorte   : " + COHORT,
        "  Dashboard : " + EXAMPLE_IP,
        "  Commande eleve (Linux/mac) :",
        "    ./scripts/install-student.sh -d " + EXAMPLE_IP + " -c " + COHORT + " -l <prenom>",
        "  Commande eleve (Windows) :",
        "    .\\scripts\\install-student.ps1 -Dashboard " + EXAMPLE_IP +
        " -Cohort " + COHORT + " -Label <prenom>",
    ])
    note(doc, "Chaque eleve doit avoir un label (-l / -Label) unique, sinon deux "
              "eleves se confondent dans la matrice.")

    doc.add_heading("6. Ouvrir ton dashboard", level=1)
    numbered(doc, "Ouvre dans ton navigateur :")
    code_block(doc, [
        "http://127.0.0.1:" + DASHBOARD_PORT + "/login",
        "http://127.0.0.1:" + DASHBOARD_PORT + "/dashboard?cohort=" + COHORT,
    ])
    numbered(doc, "Connecte-toi avec DASHBOARD_TEACHER_TOKEN (affiche par le script).")
    numbered(doc, "La matrice cohorte se met a jour en direct (flux SSE).")

    doc.add_heading("7. Reseau et securite", level=1)
    add_table(
        doc,
        ["Point", "A verifier"],
        [
            ["LAN plat", "Pas de VLAN isolant les eleves de ton poste."],
            ["Pare-feu", "Port " + DASHBOARD_PORT + " ouvert en entree sur ton poste."],
            ["CORS", "Si tous les eleves ouvrent http://127.0.0.1:" + SHOP_PORT +
                     ", la valeur par defaut de DASHBOARD_CORS_ORIGINS convient."],
            ["Jeton", "Ne partage jamais DASHBOARD_TEACHER_TOKEN avec les eleves."],
        ],
        [40, 130],
    )

    doc.add_heading("8. Arreter / nettoyer", level=1)
    code_block(doc, [
        "cd docker",
        "docker compose --env-file .env down       # arret du dashboard",
        "docker compose --env-file .env down -v    # arret + efface les donnees cohorte",
    ])
    note(doc, "down -v efface la base SQLite (toute la progression cohorte). "
              "Recupere les preuves avant si besoin.")

    annexes(doc)
    doc.save(str(PROF_OUT))
    print(f"  ecrit : {PROF_OUT}")


# ---------------------------------------------------------------------------
# Annexes communes (prerequis detailles) — appendues aux deux guides
# ---------------------------------------------------------------------------


def annexes(doc: Document) -> None:
    """Annexes prerequis : PowerShell, Docker, Git/OpenSSL, verifs et gotchas."""
    doc.add_page_break()
    doc.add_heading("Annexes — preparer le poste", level=1)
    para(doc, "A faire UNE fois par poste, avant la procedure d'installation. "
              "Si une commande echoue, voir l'annexe D (depannage).")

    # ---- A : PowerShell ---------------------------------------------------
    doc.add_heading("Annexe A — Debloquer PowerShell (Windows)", level=2)
    para(doc, "Par defaut, Windows bloque l'execution des scripts .ps1 "
              "(ExecutionPolicy = Restricted). Le script install-student.ps1 ne "
              "demarrera pas tant que ce n'est pas debloque. Aucun droit "
              "administrateur n'est requis pour les methodes ci-dessous.")
    para(doc, "Methode 1 — temporaire, pour la session courante (recommandee). "
              "Ouvre PowerShell, puis dans la MEME fenetre :")
    code_block(doc, [
        "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass",
        "# puis lance le script dans cette meme fenetre",
    ])
    para(doc, "Methode 2 — persistante pour ton compte utilisateur :")
    code_block(doc, [
        "Set-ExecutionPolicy -Scope CurrentUser -ExecutionPolicy RemoteSigned",
    ])
    para(doc, "Si le fichier a ete telecharge et est marque « bloque » par Windows :")
    code_block(doc, [
        "Unblock-File .\\scripts\\install-student.ps1",
    ])
    para(doc, "Verifier l'etat des politiques par portee :")
    code_block(doc, ["Get-ExecutionPolicy -List"])
    note(doc, "Utilise de preference PowerShell 7+ (commande pwsh) plutot que "
              "Windows PowerShell 5. Ne lance pas le script « en tant "
              "qu'administrateur » : Docker Desktop fonctionne en utilisateur normal.")

    # ---- B : Docker -------------------------------------------------------
    doc.add_heading("Annexe B — Installer Docker", level=2)

    doc.add_heading("Windows 10/11", level=3)
    numbered(doc, "Active WSL2 (PowerShell en administrateur), puis redemarre :")
    code_block(doc, ["wsl --install"])
    numbered(doc, "Installe Docker Desktop depuis docker.com (backend WSL2 par defaut).")
    numbered(doc, "Lance Docker Desktop et attends le statut « running » (icone baleine).")
    numbered(doc, "Verifie dans un terminal :")
    code_block(doc, ["docker run --rm hello-world", "docker compose version"])
    note(doc, "Erreur de virtualisation : active « Virtualization / SVM / VT-x » "
              "dans le BIOS/UEFI, et la fonctionnalite Windows « Plateforme "
              "de machine virtuelle ».")

    doc.add_heading("macOS", level=3)
    numbered(doc, "Telecharge Docker Desktop pour Mac (choisis la puce : Apple "
                  "Silicon ou Intel).")
    numbered(doc, "Glisse Docker dans Applications, lance-le, attends « running ».")
    numbered(doc, "Verifie :")
    code_block(doc, ["docker run --rm hello-world", "docker compose version"])

    doc.add_heading("Linux (Debian / Ubuntu)", level=3)
    para(doc, "Deux methodes mutuellement exclusives. Choisis l'une OU l'autre.")
    para(doc, "Methode A — paquets de la distribution (recommandee pour un TD) :")
    code_block(doc, [
        "sudo apt update",
        "sudo apt install -y docker-compose-v2",
        "sudo usermod -aG docker $USER",
        "newgrp docker   # ou se deconnecter puis se reconnecter",
        "docker compose version",
    ])
    para(doc, "Methode B — depot officiel Docker (si tu veux la derniere version) :")
    code_block(doc, [
        "# configurer le depot officiel : https://docs.docker.com/engine/install/ubuntu/",
        "sudo apt install -y docker-ce docker-ce-cli containerd.io \\",
        "    docker-buildx-plugin docker-compose-plugin",
        "sudo usermod -aG docker $USER",
        "newgrp docker",
        "docker compose version",
    ])
    note(doc, "Les deux paquets sont mutuellement exclusifs. Sur Ubuntu 25.04+, "
              "utilise la methode A (docker-compose-v2). Ne melange jamais les deux.")

    # ---- C : Git / OpenSSL ------------------------------------------------
    doc.add_heading("Annexe C — Git et OpenSSL", level=2)
    add_table(
        doc,
        ["Systeme", "Commande / source"],
        [
            ["Windows", "Git for Windows (git-scm.com) — fournit git, openssl et Git Bash"],
            ["macOS", "xcode-select --install (git) ; openssl deja present"],
            ["Linux", "sudo apt install git openssl"],
        ],
        [40, 130],
    )
    note(doc, "Le script genere les jetons avec openssl. Sous Windows, si openssl "
              "manque, il bascule sur un generateur .NET integre — aucune action requise.")

    # ---- D : Depannage prerequis -----------------------------------------
    doc.add_heading("Annexe D — Depannage des prerequis", level=2)
    add_table(
        doc,
        ["Symptome", "Cause / solution"],
        [
            ["« running scripts is disabled on this system »",
             "PowerShell bloque : voir Annexe A (Set-ExecutionPolicy -Scope Process Bypass)."],
            ["« docker : command not found » / « cannot connect to the Docker daemon »",
             "Docker Desktop n'est pas demarre, ou (Linux) groupe docker pas applique : "
             "ouvre Docker Desktop / reconnecte-toi."],
            ["Le clone git echoue ou reste bloque",
             "Proxy ou pare-feu d'entreprise : configure git (http.proxy) ou utilise "
             "un reseau ouvert pour le clone."],
            ["Le build s'arrete au telechargement d'images",
             "Proxy Docker : configure le proxy dans Docker Desktop (Settings > Resources > Proxies)."],
            ["docker compose version affiche v1.x ou commande introuvable",
             "Plugin compose v2 absent. Ubuntu/Debian standard : sudo apt install docker-compose-v2. "
             "Depot officiel Docker : sudo apt install docker-compose-plugin. "
             "Ne pas installer les deux (conflit)."],
        ],
        [60, 110],
    )


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------


def main() -> int:
    print("Generation des guides d'installation JuiceLab...")
    with tempfile.TemporaryDirectory(prefix="juicelab-docgen-") as td:
        tmp = Path(td)
        build_eleve(tmp)
        build_prof(tmp)
    print("Termine.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
